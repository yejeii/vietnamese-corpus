import tkinter as tk
from tkinter import filedialog
import tkinter.font as font
import os
import win32com.client
import docx
import shutil
import pandas as pd
from datetime import datetime
import re
import math

## 폴더 안에 순수 docx 확장자 및 xlsx 확장자 파일만 존재해야 함!!!
## .docx 확장자는 무조건 소문자로 저장되어야 함!!


inappropriate_files = []

class UnCorrectFileNameException(Exception):
    def __init__(self, message):
        super().__init__(message)

class DataReadException(Exception):
    def __init__(self, message):
        super().__init__(message)

def validate_docx_filename(filename):
    pattern =r'\d{2}_CW\d{3}_\d{8}_\d{4}\.docx'
    
    if re.match(pattern, filename):
        return True
    else:
        raise UnCorrectFileNameException(f"Modify the file name '{filename}' to match the file name format.")

def validate_xlsx_filename(filename):
    # 파일명에 해당하는 정규 표현식 패턴
    pattern = r'\d{8}_CW\d{3}_\d{1,4}.xlsx'

    if re.match(pattern, filename):
        return True
    else:
        raise UnCorrectFileNameException(f"Modify the file name '{filename}' to match the file name format.")

def checkValues(index, row, worker_id, job_ymd):
    fill_in = False
    datas = [str(row.iloc[i]) for i in [0, 1, 2, 3, 4, 5, 6, 7, 8, 9, 10, 11, 12, 13, 14]]
    # print(datas)
    
    if any(data == 'nan' for data in datas):
        raise DataReadException(f"Make sure the {index+2} row contains all your input values!")
    else:
        fill_in = True
        
    if fill_in:
        if str(row.iloc[1]) == worker_id and str(row.iloc[2]) == job_ymd and len(str(row.iloc[3])) == 8:
            return True
        else:
            raise DataReadException(f"Check the worker name or work date and publication date in the {index+2} line at excel file!")

def get_file_info(file_path):
    file_size = os.path.getsize(file_path)                  # 파일 크기를 바이트 단위로 가져오기 
    file_creation_time = os.path.getctime(file_path)        # 파일 생성일자를 가져오기
    file_modification_time = os.path.getmtime(file_path)    # 파일 최종 수정일자를 가져오기
    file_name = os.path.basename(file_path)                 # 파일명 추출
    
    return file_name, file_size, file_creation_time, file_modification_time

def count_words_doc(file_path):
    word_app = win32com.client.Dispatch("Word.Application")
    doc = word_app.Documents.Open(file_path)
    word_count = doc.Words.Count
    paragraphs = doc.Content.Text.split('\n')

    doc.Close()
    word_app.Quit()
    
    for paragraph in paragraphs:
        if paragraph:
            paragraph = str(paragraph).replace("'", "''")
    
    return word_count, paragraph

def count_words_docx(file_path):
    doc = docx.Document(file_path)

    total_words = 0
    
    for paragraph in doc.paragraphs:
        total_words += len(paragraph.text.split())   # 문서 내 단어 수 세기 
        text = paragraph.text.strip()
        if text:
            text = str(text).replace("'", "''")

    return total_words, text

def format_time(timestamp):
    dt_object = datetime.fromtimestamp(timestamp)               # 에포크 시간을 datetime 객체로 변환
    formatted_time = dt_object.strftime('%Y-%m-%d %H:%M:%S')    # 원하는 형식으로 시간을 포맷팅
    return formatted_time

def make_inappropriate_folder(root_path):
    inappropriate_folder = root_path+'inappropriate_folder'
    if not os.path.exists(inappropriate_folder):
        os.makedirs(inappropriate_folder)
    return inappropriate_folder

def move_files_to_inappropriate_folder(files, inappropriate_folder):
    for move_file in files:
        shutil.move(move_file, inappropriate_folder)
        print(f'{move_file} move to "inappropriate_folder"')

def browse_folder():
    folder_path = filedialog.askdirectory()  # 폴더 선택 대화 상자 열기
    if folder_path:
        entry_path.delete(0, tk.END)  # 입력 필드 초기화
        entry_path.insert(0, folder_path)  # 선택한 경로 입력 필드에 설정

def register_path():
    
    first_value = ''
    second_value = ''
    third_value = ''
    
    is_error = False
    
    path = entry_path.get()  # 입력 필드에서 경로 가져오기
    
    print(path)
    if os.path.exists(path+'/inappropriate_folder'):
        os.rmdir(path+'/inappropriate_folder')
    
    try:
        file_path = ''
        
        if path:
            
            folder = os.path.basename(path)  # 마지막 폴더명 가져오기
            
            parts = folder.split('_')  # 구분자로 분리..
            if len(parts) == 3:                 # 3 개로 분리 되어 있는경우    20230821_CW002_241
                first_value = parts[0]
                second_value = parts[1]
                third_value = parts[2]

            # root_path = 'C:/Work/과제98_베트남어말뭉치데이터/00_수집자료/00_완료_추가_0831/CW045/20230831_CW045_232'
            root_path = path + '/'
            file_list = os.listdir(root_path)
            file_list = [file for file in file_list if not file.startswith('~$') and file != '.DS_Store']
            docx_cnt = len(file_list) - 1   # 작업한 docx 개수
            print(docx_cnt)
            
            if int(third_value) != int(docx_cnt):
                raise DataReadException(f"The number of docx files and the number of files written in the folder name are different.!!!")
            else:
                for file in file_list:
                    if not (file.startswith('~$') or file.startswith('.DS_Store')):
                        file_path = root_path+file

                        # 파일 정보 가져오기
                        file_name, file_size, file_creation_time, file_modification_time = get_file_info(file_path)

                              
                        # 단어수 업데이트 할 때
                        if file_path.lower().endswith(".docx"):
                            print(f"file_name : {file_name} start to read...")
                            try:
                                # docx 파일명 형식 확인
                                if validate_docx_filename(file_name):
                                    doc = docx.Document(file_path)
                                    total_words = 0
                                    
                                    for paragraph in doc.paragraphs:
                                        total_words += len(paragraph.text.split())   # 문서 내 단어 수 세기 

                                    if total_words == 0:
                                        raise DataReadException(f'Error occured in {file_name} : Can not read Data. Check the file.')
                            except Exception as e:
                                print(f'Error occured in {file_name} : {e}')
                                inappropriate_files.append(file_path) 
                        
                        elif file_path.lower().endswith(".doc"):
                            print(".doc is unsupported file format.")
                            inappropriate_files.append(file_path)
                        
                        elif file_path.lower().endswith(".xlsx"):
                            xlsx_file = os.path.basename(file_path)
                            
                            # 1. 파일명 확인
                            if validate_xlsx_filename(xlsx_file):
                                
                                job_ymd = ''
                                worker_id = ''
                                work_cnt = ''
                                
                                parts = xlsx_file.split('_')  # 구분자로 분리..
                                if len(parts) == 3:                
                                    job_ymd = parts[0]
                                    worker_id = parts[1]
                                    work_cnt = parts[2]
                                
                                # 2. 데이터 read 확인
                                datas = []
                                read_flag = False
                                read_success = False
                                
                                data_frame = pd.read_excel(file_path, engine='openpyxl')
                                datas.append(data_frame)
                                
                                # 5번째 행 첫번째 열의 값 확인
                                # print(datas[0].iloc[3])
                                fifth_row_first_col = datas[0].iloc[3, 0]
                                
                                if str(fifth_row_first_col) == '1':
                                    read_flag = True    
                                else:
                                    raise DataReadException(f'Double check that your data starts from row 5!')
                                
                                if read_flag:
                                    for df in datas:
                                        for index, row in df.iloc[3:].iterrows():   # 5행부터 read
                                            
                                            # 3-1, 3-3, 3-4. 데이터 유무, 작업자코드, 작업날짜 비교
                                            if checkValues(index, row, worker_id, job_ymd):
                                                read_success = True

                                # 데이터 성공적으로 읽은 경우
                                if read_success:
                                    print(f'File {xlsx_file} read successfully.')    
                        
                        else:
                            print(f"{file} : Unsupported file format.")   
                        
    except UnCorrectFileNameException as e:
        print(f'Operator Modifications : {e}')
        inappropriate_files.append(file_path)
    except DataReadException as e:
        print(f'Operator Modifications :{e}')
        inappropriate_files.append(file_path)
    except Exception as e:
        print(f'Error occured in {file_name} : {e}')
        inappropriate_files.append(file_path)       
    
    finally:
        inappropriate_folder = ''
        if len(inappropriate_files) > 0:
            # 부적합 폴더 생성, 부적합 파일 이동
            inappropriate_folder = make_inappropriate_folder(root_path)
            move_files_to_inappropriate_folder(inappropriate_files, inappropriate_folder)
            print("Check the 'inappropriate_folder' folder and revise files!!")
        else:
            if os.path.exists(inappropriate_folder):
                os.rmdir(inappropriate_folder)
            print("The folder has been successfully READ!!")

# 메인 창 생성
root = tk.Tk()

root.customFont = font.Font(family='Helvetica', size=15)
root.title("Data Upload Program")

# 라벨 생성
label = tk.Label(root, text="Choose the path that contains the materials you completed today! (Ex. : 20230816_CW004_260)", font=root.customFont)
label.pack()

# 입력 필드 생성
entry_path = tk.Entry(root, width=120)
entry_path.pack()

# "찾아보기" 버튼 생성
browse_button = tk.Button(root, text="Search", command=browse_folder,font=13)
browse_button.pack()

# "실행" 버튼 생성
register_button = tk.Button(root, text="execute", command=register_path, font=13)
register_button.pack()

# 텍스트 위젯 생성
text = "[Inappropriate folders] that are created must be corrected.!!!"
label2 = tk.Label(root, text=text, 
                  font=root.customFont, 
                  foreground="red",
                  width=80, height=5)
label2.pack()

# 프로그램 실행
root.mainloop()
