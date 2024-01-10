import sys

import docx
import os
from datetime import datetime

import pandas as pd
import pymysql
from pymysql import MySQLError

####################################################################################################################
# MySQL 연결 정보
# mysql_host = '172.30.1.36'
# mysql_port = 13333
mysql_host = 'localhost'
mysql_port = 3307
mysql_user = 'vnc'
mysql_password = 'vnc'
mysql_database = 'vncsim'
db = pymysql.connect(host=mysql_host,port=mysql_port,user=mysql_user,passwd=mysql_password,db=mysql_database)
#####################################################################################################################

root_path = r'D:/YJ/'
JOB_YMD = '20230905'
WORKER_ID = ''

# 디비 TB명
excel_tb = 'load_vnc_org_lst_sim_'+JOB_YMD
file_tb = 'load_vnc_file_list_sim_'+JOB_YMD
docx_tb = 'load_vnc_file_docx_sim_'+JOB_YMD

# 테이블 생성하는 쿼리문 결과
CREATE_TABLE = False  # 기본적으로 fALSE로 초기화

# 벌크 변수
chunk_size = 300

def get_file_info(file_path):
    file_size = os.path.getsize(file_path)                  # 파일 크기를 바이트 단위로 가져오기 
    file_creation_time = os.path.getctime(file_path)        # 파일 생성일자를 가져오기
    file_modification_time = os.path.getmtime(file_path)    # 파일 최종 수정일자를 가져오기
    file_name, file_extension = os.path.splitext(os.path.basename(file_path))                 # 확장명 뺀 파일명 추출
    
    return file_name, file_size, file_creation_time, file_modification_time

def format_time(timestamp):
    dt_object = datetime.fromtimestamp(timestamp)               # 에포크 시간을 datetime 객체로 변환
    formatted_time = dt_object.strftime('%Y-%m-%d %H:%M:%S')    # 원하는 형식으로 시간을 포맷팅
    
    return formatted_time

def count_words_docx(file_path):
    total_words = 0
    doc = docx.Document(file_path)

    for paragraph in doc.paragraphs:
        total_words += len(paragraph.text.split())   # 문서 내 단어 수 세기 

    return total_words

def create_table(db, cursor):
    # DDL - auto commit. db.commit() X
    try:         
        create_query = f'CREATE TABLE IF NOT EXISTS {excel_tb} LIKE vnc.load_vnc_org_lst_sim;'
        cursor.execute(create_query)    # 쿼리 실행

        create_query = f'CREATE TABLE IF NOT EXISTS {file_tb} LIKE vnc.load_vnc_file_list_sim;'
        cursor.execute(create_query)

        create_query = f'CREATE TABLE IF NOT EXISTS {docx_tb} LIKE vnc.load_vnc_file_docx_sim;'
        cursor.execute(create_query)
    
        # create_idx_q = f'CREATE INDEX IF NOT EXISTS {docx_tb}_FILE_TXT_IDX USING BTREE ON {docx_tb} (FILE_TXT);'
        # cursor.execute(create_idx_q)    # 쿼리 실행
        # db.commit() # 변경사항 DB에 반영
        
        print("Create succeeded.")
        CREATE_TABLE = True
        
    except Exception as e:
        print(f"Error at creating tables : {e}")
        sql = ''
        sql += f'DROP TABLE IF EXISTS {excel_tb};'
        cursor.execute(sql)
        
        sql = f'DROP TABLE IF EXISTS {file_tb};'
        cursor.execute(sql)
        
        sql = f'DROP TABLE IF EXISTS {docx_tb};'
        cursor.execute(sql)

        # db.commit()

    finally:
        return CREATE_TABLE

######################################################################################################################################


try:
    #  MySQL 데이터베이스에 연결
    db = pymysql.connect(host=mysql_host,port=mysql_port,user=mysql_user,passwd=mysql_password,db=mysql_database)   
    cursor = db.cursor(pymysql.cursors.DictCursor)  # 데이터베이스와 상호작용하기 위한 커서(cursor)를 생성
    print("Connected to MySQL database")
        
    # 테이블 CREATE
    if create_table(db=db, cursor=cursor):
        
        job_ymd_path = os.path.join(root_path, JOB_YMD)
        worker_folders = os.listdir(job_ymd_path)
        for worker_folder in worker_folders:
            WORKER_ID = worker_folder.split('_')[1]
            
            if len(WORKER_ID) > 5 :
                # 작업자 폴더 길이가 5를 초과하면 해당 파일 처리를 건너뜀.(자동화로 처리된 파일)
                continue
            
            print(JOB_YMD, WORKER_ID + ' 처리 시작')
            
            # 작업자의 .docx, .xlsx에 접근
            worker_folder_pth = os.path.join(job_ymd_path, worker_folder)
            files = os.listdir(worker_folder_pth)
            files = [file for file in files if not file.startswith('~$') and (file.endswith('.docx') or file.endswith('.xlsx'))]

            WORK_CNT = len(files) - 1    # docx 파일 개수
            
            
            # 데이터 담는 배열변수
            excel_bulk_insert_datas = []
            file_bulk_insert_datas = []
            docx_bulk_insert_datas = []
            
            for file in files:
                file_pth = os.path.join(worker_folder_pth, file)
                
                # .docx 처리
                if file_pth.lower().endswith(".docx"):
                    
                    # {file_db} 
                    # 파일 정보 가져오기
                    FILE_NAME, FILE_SIZE, file_creation_time, file_modification_time = get_file_info(file_pth)

                    # 시간을 원하는 형식으로 변환
                    FILE_CREATETIME = format_time(file_creation_time)
                    FILE_CHANG_TIME = format_time(file_modification_time)
                    
                    # {docx_tb} 처리
                    doc = docx.Document(file_pth)
                    SEQ = 1
                    WORD_COUNT = 0
                    # docx_bulk_insert_datas = []
                    
                    for paragraph in doc.paragraphs:
                        WORD_COUNT += len(paragraph.text.split())   # 문서 내 단어 수 세기 

                        texts = paragraph.text.split('\n')
                        texts = [txts for txts in texts if not txts == '' ]           # 빈공백으로만 있는 텍스트 제거
                        texts = [txts for txts in texts if not txts.strip() == '' ]   # 앞뒤 공백 자른 txts가 ''인 것 제거
                        
                        # '\n' 별로 저장할 때           
                        if texts:
                            for text in texts:
                                text = text.replace("'", "''")         # 디비 인서트를 위한 처리
                                text = text.replace("\u00A0", " ")    # strip()이 걸러내지 못하는 &nbsp 제거
                                text = text.replace("\uFEFF", " ")    # HTML 코드(&#65279;) 제거
                                text = text.replace('\t', ' ')
                                text = text.replace('\xa0', ' ')
                                text = text.strip()     # 앞 뒤 공백 제거
                                # print(f'{SEQ} : {text} \n')
                                
                                # 벌크에 넣기
                                docx_bulk_insert_datas.append((FILE_NAME, SEQ, text))

                                # text line_seq 증가
                                SEQ += 1
                    
                    file_bulk_insert_datas.append((JOB_YMD, WORKER_ID, WORK_CNT, FILE_NAME, FILE_SIZE, FILE_CREATETIME, FILE_CHANG_TIME, WORD_COUNT))
                    print(f'{FILE_NAME} fin')

                elif file_pth.endswith(".xlsx"):
                    
                    # 엑셀에서 데이터 GET
                    # bulk_insert_datas = []
                    datas = []
                    read_flag = False

                    data_frame = pd.read_excel(file_pth, engine='openpyxl')
                    datas.append(data_frame)
                    
                    for df in datas:
                        for index, row in df.iterrows():
                        
                            if(str(row.iloc[0]) == '1'):
                                read_flag = True
                            
                            if(str(row.iloc[0]) == ''):
                                read_flag = False
                            
                            if(read_flag):
                                SEQ = int(row.iloc[0])
                                PUB_YMD = str(row.iloc[3]).strip()
                                WTR_KR = str(row.iloc[4]).replace("'", "''").strip()
                                WTR_VN = str(row.iloc[5]).replace("'", "''").strip()
                                CTR_KR = str(row.iloc[6]).replace("'", "''").strip()
                                CTR_VN = str(row.iloc[7]).replace("'", "''").strip()
                                ORG_TYP = str(row.iloc[8]).strip()
                                DAT_SRC = str(row.iloc[9]).replace("'", "''").strip()   # 작은따옴표를 두 번 사용하여 이스케이프
                                DAT_TYP = str(row.iloc[10]).strip()
                                TOPIC_KR = str(row.iloc[11]).strip()
                                TOPIC_VN = str(row.iloc[12]).strip()
                                TITLE = str(row.iloc[13]).replace("'", "''").strip()  # 작은따옴표를 두 번 사용하여 이스케이프
                                COL_PK = str(row.iloc[14]).split('.')[0]
                                DOC_STYLE = '구어체' if '구어체' in str(row.iloc[15]).strip() else '문어체'
                                # DOC_STYLE = '구어체'
                                # print(SEQ, COL_PK, PUB_YMD, TITLE)
                                                            
                                # {excel_tb}에 인서트
                                excel_bulk_insert_datas.append((COL_PK, SEQ, WORKER_ID, JOB_YMD, PUB_YMD, WTR_KR, 
                                                                WTR_VN, CTR_KR, CTR_VN, ORG_TYP, DAT_SRC, DAT_TYP, 
                                                                TOPIC_KR, TOPIC_VN, TITLE, COL_PK, DOC_STYLE))
                                
                    print(f'{file} fin')
                    
            # 작업자별로 인서트    
            sql = ''
            sql = f'SAVEPOINT {WORKER_ID}'
            cursor.execute(sql)
            
            # {file_tb} 인서트
            sql = ""
            sql = sql + f" INSERT INTO {file_tb} (workymd, worker_id, workcnt, file_name, file_size, file_createtime, file_chang_time, word_count)"
            sql = sql + " VALUES (%s, %s, %s, %s, %s, %s, %s, %s)"
            cursor.executemany(sql, file_bulk_insert_datas)
            # db.commit()
            
            # {docx_tb} 300개씩 인서트
            for i in range(0, len(docx_bulk_insert_datas), 300):
                batch = docx_bulk_insert_datas[i:i+300]
                
                sql = ""
                sql = sql + f" INSERT INTO {docx_tb} (FILE_NAME, SEQ, FILE_TXT)"
                sql = sql + " VALUES (%s, %s, %s)"
                cursor.executemany(sql, batch)
                # db.commit()   
                # print(f'{len(batch)}')
            
            # {excel_tb} 인서트
            sql = ""
            sql = sql + f" INSERT INTO {excel_tb} "
            sql = sql + " (COL_PK, SEQ, WORKER_ID, JOB_YMD, PUB_YMD, WTR_KR, WTR_VN, CTR_KR, CTR_VN, ORG_TYP, DAT_SRC, DAT_TYP, TOPIC_KR, TOPIC_VN, TITLE, FILE_NAME, DOC_STYLE) "
            sql = sql + " VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s) "
            cursor.executemany(sql, excel_bulk_insert_datas)
            # db.commit()
            
            # 3개의 테이블 모두 execute 후 db 영구반영         
            db.commit()    
            
            print(f"{WORKER_ID} FIN")        
    else:
        print("테이블 생성 실패")

except Exception as e:
    print(f"Error at preceeding {WORKER_ID} files")
    print(f"Error: {e}")
    
    sql = ''
    sql = f'rollback to {WORKER_ID}'
    cursor.execute(sql)
    db.commit()
    
    sys.exit()

finally:
    if cursor:
        cursor.close()
    if db:
        db.close()
        print("Disconnected from MySQL database \n")                    